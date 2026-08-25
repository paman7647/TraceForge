import csv
import datetime
import hashlib
import json
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from traceforge.case import Case

RE_IP = re.compile(r"\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b")
RE_EMAIL = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")

def sanitize_csv_cell(val: Any) -> str:
    """Defends against CSV / Spreadsheet formula injection attacks."""
    s = str(val) if val is not None else ""
    if not s:
        return s
    if s[0] in ("=", "+", "-", "@", "\t", "\r"):
        return "'" + s
    return s

def redact_text(text: str) -> str:
    """Masks sensitive IP addresses and email addresses in generated reports."""
    def mask_ip(m):
        raw = m.group(0)
        h = hashlib.sha256(raw.encode()).hexdigest()[:4].upper()
        return f"REDACTED_IP_{h}"

    def mask_email(m):
        raw = m.group(0)
        h = hashlib.sha256(raw.encode()).hexdigest()[:4].upper()
        return f"redacted_user_{h}@redacted.local"

    res = RE_IP.sub(mask_ip, text)
    res = RE_EMAIL.sub(mask_email, res)
    return res

class CaseExporter:
    """Generates multi-format forensic reports and threat intelligence datasets."""

    def __init__(self, case: Case, redact: bool = False):
        self.case = case
        self.redact = redact
        self.data = case.data.copy()

    def _filter_str(self, s: str) -> str:
        return redact_text(s) if self.redact else s

    def export_all(self, out_dir: Optional[Union[str, Path]] = None) -> Dict[str, Path]:
        dest = Path(out_dir) if out_dir else self.case.exports_dir
        if self.redact:
            dest = dest / "redacted"
        dest.mkdir(parents=True, exist_ok=True)

        results = {
            "markdown": self.export_markdown(dest / f"{self.case.case_id}.md"),
            "html": self.export_html(dest / f"{self.case.case_id}.html"),
            "json": self.export_json(dest / f"{self.case.case_id}.json"),
            "jsonl_timeline": self.export_jsonl_timeline(dest / f"{self.case.case_id}_timeline.jsonl"),
            "csv_iocs": self.export_csv_iocs(dest / "iocs.csv"),
            "csv_evidence": self.export_csv_evidence(dest / "evidence.csv"),
            "csv_findings": self.export_csv_findings(dest / "findings.csv"),
            "csv_timeline": self.export_csv_timeline(dest / "timeline.csv"),
            "stix": self.export_stix(dest / f"{self.case.case_id}_stix21.json"),
            "misp": self.export_misp(dest / f"{self.case.case_id}_misp.json"),
            "geojson": self.export_geojson(dest / f"{self.case.case_id}.geojson"),
            "kml": self.export_kml(dest / f"{self.case.case_id}.kml"),
        }

        # Optional XLSX / DOCX if libraries are installed
        try:
            results["xlsx"] = self.export_xlsx(dest / f"{self.case.case_id}.xlsx")
        except Exception:
            pass

        try:
            results["docx"] = self.export_docx(dest / f"{self.case.case_id}.docx")
        except Exception:
            pass

        return results

    def export_markdown(self, out_file: Path) -> Path:
        summary = self.case.get_summary()
        lines = [
            f"# Forensic Investigation Report: {summary['case_id']}",
            f"\n**Case Title**: {self._filter_str(summary['case_name'])}",
            f"**Lead Analyst**: {summary['analyst']}",
            f"**Created At**: {summary['created_at']}",
            f"**Status**: {summary['status'].upper()}\n",
            "---",
            "\n## Executive Summary\n",
            f"- **Total Ingested Evidence**: {summary['total_evidence']}",
            f"- **Findings**: {summary['total_findings']} ({summary['high_severity_findings']} High/Critical)",
            f"- **Threat Indicators (IOCs)**: {summary['total_iocs']}",
            f"- **Timeline Milestones**: {summary['total_timeline_events']}\n",
            "---",
            "\n## Findings & Observations\n",
        ]

        findings = self.data.get("findings", [])
        if not findings:
            lines.append("_No findings recorded._\n")
        else:
            for f in findings:
                lines.append(f"### [{f.get('severity','-').upper()}] {self._filter_str(f.get('title','Untitled'))}")
                lines.append(f"**ID**: `{f.get('id','-')}` | **Category**: {f.get('category','-')} | **Status**: {f.get('status','-')}")
                lines.append(f"\n{self._filter_str(f.get('description',''))}\n")

        lines.append("---")
        lines.append("\n## Indicators of Compromise (IOCs)\n")
        iocs = self.data.get("iocs", [])
        if not iocs:
            lines.append("_No IOCs registered._\n")
        else:
            lines.append("| ID | Type | Observable Value | Confidence | Context |")
            lines.append("|---|---|---|---|---|")
            for i in iocs:
                val = self._filter_str(i.get('value','-'))
                lines.append(f"| `{i.get('id','-')}` | `{i.get('type','-')}` | `{val}` | {i.get('confidence','-')} | {self._filter_str(i.get('context','-'))} |")
            lines.append("")

        lines.append("---")
        lines.append("\n## Chronological Timeline\n")
        events = self.data.get("timeline", [])
        if not events:
            lines.append("_No timeline events recorded._\n")
        else:
            lines.append("| Timestamp | Source | Title / Event | Severity |")
            lines.append("|---|---|---|---|")
            for e in events:
                lines.append(f"| {e.get('timestamp','-')} | {e.get('source','-')} | {self._filter_str(e.get('title','-'))} | `{e.get('severity','-')}` |")
            lines.append("")

        lines.append("---")
        lines.append("\n## Evidence Inventory & Cryptographic Hashes\n")
        evidence = self.data.get("evidence", [])
        if not evidence:
            lines.append("_No evidence files attached._\n")
        else:
            lines.append("| ID | File Name | Size (Bytes) | SHA-256 Digest | Description |")
            lines.append("|---|---|---|---|---|")
            for ev in evidence:
                lines.append(f"| `{ev.get('id','-')}` | {ev.get('filename','-')} | {ev.get('size_bytes',0)} | `{ev.get('sha256','-')}` | {self._filter_str(ev.get('description','-'))} |")
            lines.append("")

        content = "\n".join(lines)
        out_file.parent.mkdir(parents=True, exist_ok=True)
        with open(out_file, "w", encoding="utf-8") as fp:
            fp.write(content)
        return out_file

    def export_html(self, out_file: Path) -> Path:
        summary = self.case.get_summary()
        title = f"TraceForge Report — {summary['case_id']}"

        findings_rows = []
        for f in self.data.get("findings", []):
            sev = f.get("severity", "info").lower()
            findings_rows.append(f"<tr><td><span class='badge badge-{sev}'>{sev.upper()}</span></td><td><strong>{self._filter_str(f.get('title',''))}</strong></td><td>{f.get('category','')}</td><td>{self._filter_str(f.get('description',''))}</td></tr>")

        ioc_rows = []
        for i in self.data.get("iocs", []):
            ioc_rows.append(f"<tr><td><code>{i.get('id','')}</code></td><td>{i.get('type','')}</td><td><code>{self._filter_str(i.get('value',''))}</code></td><td>{i.get('confidence','')}</td><td>{self._filter_str(i.get('context',''))}</td></tr>")

        evidence_rows = []
        for ev in self.data.get("evidence", []):
            evidence_rows.append(f"<tr><td><code>{ev.get('id','')}</code></td><td>{ev.get('filename','')}</td><td>{ev.get('size_bytes',0)}</td><td><code class='hash'>{ev.get('sha256','')}</code></td><td>{self._filter_str(ev.get('description',''))}</td></tr>")

        timeline_rows = []
        for e in self.data.get("timeline", []):
            timeline_rows.append(f"<tr><td>{e.get('timestamp','')}</td><td>{e.get('source','')}</td><td><strong>{self._filter_str(e.get('title',''))}</strong><br>{self._filter_str(e.get('description',''))}</td><td><code>{e.get('severity','info')}</code></td></tr>")

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<style>
body {{ background: #0f172a; color: #f8fafc; font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, monospace; margin: 0; padding: 28px; }}
h1 {{ color: #38bdf8; font-size: 22px; margin-bottom: 6px; }}
h2 {{ color: #94a3b8; font-size: 16px; border-bottom: 1px solid #334155; padding-bottom: 8px; margin-top: 28px; text-transform: uppercase; letter-spacing: 0.5px; }}
.meta {{ color: #94a3b8; font-size: 13px; margin-bottom: 20px; }}
.stats {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(180px, 1fr)); gap: 16px; margin-bottom: 24px; }}
.card {{ background: #1e293b; border: 1px solid #334155; border-radius: 6px; padding: 14px; }}
.card-label {{ color: #94a3b8; font-size: 11px; text-transform: uppercase; font-weight: 600; }}
.card-val {{ font-size: 22px; font-weight: bold; color: #38bdf8; margin-top: 4px; }}
table {{ width: 100%; border-collapse: collapse; background: #1e293b; border-radius: 6px; overflow: hidden; margin-top: 12px; font-size: 13px; }}
th {{ background: #0b132b; color: #94a3b8; text-align: left; padding: 10px 14px; font-size: 11px; text-transform: uppercase; }}
td {{ padding: 10px 14px; border-bottom: 1px solid #334155; vertical-align: top; }}
tr:hover {{ background: #273549; }}
.badge {{ display: inline-block; padding: 2px 6px; border-radius: 4px; font-size: 11px; font-weight: bold; }}
.badge-critical {{ background: #ef4444; color: #fff; }}
.badge-high {{ background: #f97316; color: #fff; }}
.badge-medium {{ background: #eab308; color: #000; }}
.badge-low {{ background: #3b82f6; color: #fff; }}
.badge-info {{ background: #64748b; color: #fff; }}
code.hash {{ font-size: 11px; color: #cbd5e1; word-break: break-all; }}
</style>
</head>
<body>
<h1>{title}</h1>
<div class="meta">Case: {summary['case_name']} | Analyst: {summary['analyst']} | Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</div>

<div class="stats">
  <div class="card"><div class="card-label">Evidence Files</div><div class="card-val">{summary['total_evidence']}</div></div>
  <div class="card"><div class="card-label">Total Findings</div><div class="card-val">{summary['total_findings']}</div></div>
  <div class="card"><div class="card-label">High / Critical</div><div class="card-val">{summary['high_severity_findings']}</div></div>
  <div class="card"><div class="card-label">Observables (IOCs)</div><div class="card-val">{summary['total_iocs']}</div></div>
  <div class="card"><div class="card-label">Timeline Events</div><div class="card-val">{summary['total_timeline_events']}</div></div>
</div>

<h2>Findings & Discoveries</h2>
<table>
<thead><tr><th>Severity</th><th>Title</th><th>Category</th><th>Details</th></tr></thead>
<tbody>{"".join(findings_rows) if findings_rows else "<tr><td colspan='4'>No findings registered.</td></tr>"}</tbody>
</table>

<h2>Threat Observables (IOCs)</h2>
<table>
<thead><tr><th>ID</th><th>Type</th><th>Value</th><th>Confidence</th><th>Context</th></tr></thead>
<tbody>{"".join(ioc_rows) if ioc_rows else "<tr><td colspan='5'>No threat observables recorded.</td></tr>"}</tbody>
</table>

<h2>Evidence Inventory</h2>
<table>
<thead><tr><th>ID</th><th>Filename</th><th>Size (Bytes)</th><th>SHA-256 Digest</th><th>Description</th></tr></thead>
<tbody>{"".join(evidence_rows) if evidence_rows else "<tr><td colspan='5'>No evidence files recorded.</td></tr>"}</tbody>
</table>

<h2>Timeline Events</h2>
<table>
<thead><tr><th>Timestamp (UTC)</th><th>Source</th><th>Event</th><th>Severity</th></tr></thead>
<tbody>{"".join(timeline_rows) if timeline_rows else "<tr><td colspan='4'>No timeline events recorded.</td></tr>"}</tbody>
</table>

</body>
</html>"""
        out_file.parent.mkdir(parents=True, exist_ok=True)
        with open(out_file, "w", encoding="utf-8") as fp:
            fp.write(html)
        return out_file

    def export_json(self, out_file: Path) -> Path:
        out_file.parent.mkdir(parents=True, exist_ok=True)
        data = self.data.copy()
        if self.redact:
            data = json.loads(redact_text(json.dumps(data)))
        with open(out_file, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2)
        return out_file

    def export_jsonl_timeline(self, out_file: Path) -> Path:
        out_file.parent.mkdir(parents=True, exist_ok=True)
        events = self.data.get("timeline", [])
        with open(out_file, "w", encoding="utf-8") as f:
            for e in events:
                rec = e.copy()
                if self.redact:
                    rec = json.loads(redact_text(json.dumps(rec)))
                f.write(json.dumps(rec) + "\n")
        return out_file

    def export_csv_iocs(self, out_file: Path) -> Path:
        out_file.parent.mkdir(parents=True, exist_ok=True)
        headers = ["id", "type", "value", "confidence", "context", "first_seen", "last_seen"]
        with open(out_file, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow([sanitize_csv_cell(h) for h in headers])
            for i in self.data.get("iocs", []):
                val = self._filter_str(i.get("value", ""))
                ctx = self._filter_str(i.get("context", ""))
                writer.writerow([
                    sanitize_csv_cell(i.get("id", "")),
                    sanitize_csv_cell(i.get("type", "")),
                    sanitize_csv_cell(val),
                    sanitize_csv_cell(i.get("confidence", "")),
                    sanitize_csv_cell(ctx),
                    sanitize_csv_cell(i.get("first_seen", "")),
                    sanitize_csv_cell(i.get("last_seen", "")),
                ])
        return out_file

    def export_csv_evidence(self, out_file: Path) -> Path:
        out_file.parent.mkdir(parents=True, exist_ok=True)
        headers = ["id", "filename", "size_bytes", "sha256", "md5", "description", "acquired_at"]
        with open(out_file, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow([sanitize_csv_cell(h) for h in headers])
            for ev in self.data.get("evidence", []):
                writer.writerow([
                    sanitize_csv_cell(ev.get("id", "")),
                    sanitize_csv_cell(ev.get("filename", "")),
                    sanitize_csv_cell(ev.get("size_bytes", 0)),
                    sanitize_csv_cell(ev.get("sha256", "")),
                    sanitize_csv_cell(ev.get("md5", "")),
                    sanitize_csv_cell(self._filter_str(ev.get("description", ""))),
                    sanitize_csv_cell(ev.get("acquired_at", "")),
                ])
        return out_file

    def export_csv_findings(self, out_file: Path) -> Path:
        out_file.parent.mkdir(parents=True, exist_ok=True)
        headers = ["id", "title", "severity", "status", "category", "description", "created_at"]
        with open(out_file, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow([sanitize_csv_cell(h) for h in headers])
            for fi in self.data.get("findings", []):
                writer.writerow([
                    sanitize_csv_cell(fi.get("id", "")),
                    sanitize_csv_cell(self._filter_str(fi.get("title", ""))),
                    sanitize_csv_cell(fi.get("severity", "")),
                    sanitize_csv_cell(fi.get("status", "")),
                    sanitize_csv_cell(fi.get("category", "")),
                    sanitize_csv_cell(self._filter_str(fi.get("description", ""))),
                    sanitize_csv_cell(fi.get("created_at", "")),
                ])
        return out_file

    def export_csv_timeline(self, out_file: Path) -> Path:
        out_file.parent.mkdir(parents=True, exist_ok=True)
        headers = ["id", "timestamp", "source", "title", "severity", "description"]
        with open(out_file, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow([sanitize_csv_cell(h) for h in headers])
            for e in self.data.get("timeline", []):
                writer.writerow([
                    sanitize_csv_cell(e.get("id", "")),
                    sanitize_csv_cell(e.get("timestamp", "")),
                    sanitize_csv_cell(e.get("source", "")),
                    sanitize_csv_cell(self._filter_str(e.get("title", ""))),
                    sanitize_csv_cell(e.get("severity", "")),
                    sanitize_csv_cell(self._filter_str(e.get("description", ""))),
                ])
        return out_file

    def export_stix(self, out_file: Path) -> Path:
        out_file.parent.mkdir(parents=True, exist_ok=True)
        stix_bundle = {
            "type": "bundle",
            "id": f"bundle--{hashlib.sha256(self.case.case_id.encode()).hexdigest()[:32]}",
            "objects": []
        }
        for i in self.data.get("iocs", []):
            stix_bundle["objects"].append({
                "type": "indicator",
                "id": f"indicator--{hashlib.sha256(i.get('id','').encode()).hexdigest()[:32]}",
                "created": i.get("first_seen", datetime.datetime.now(datetime.timezone.utc).isoformat()),
                "name": self._filter_str(i.get("value", "")),
                "pattern": f"[{i.get('type','domain')}:value = '{self._filter_str(i.get('value',''))}']",
                "pattern_type": "stix",
            })
        with open(out_file, "w", encoding="utf-8") as f:
            json.dump(stix_bundle, f, indent=2)
        return out_file

    def export_misp(self, out_file: Path) -> Path:
        out_file.parent.mkdir(parents=True, exist_ok=True)
        misp_event = {
            "Event": {
                "info": self._filter_str(self.data.get("case_name", self.case.case_id)),
                "date": datetime.datetime.now().strftime("%Y-%m-%d"),
                "Attribute": []
            }
        }
        for i in self.data.get("iocs", []):
            misp_event["Event"]["Attribute"].append({
                "type": i.get("type", "domain"),
                "value": self._filter_str(i.get("value", "")),
                "comment": self._filter_str(i.get("context", "")),
            })
        with open(out_file, "w", encoding="utf-8") as f:
            json.dump(misp_event, f, indent=2)
        return out_file

    def export_geojson(self, out_file: Path) -> Path:
        out_file.parent.mkdir(parents=True, exist_ok=True)
        geojson = {
            "type": "FeatureCollection",
            "features": []
        }
        with open(out_file, "w", encoding="utf-8") as f:
            json.dump(geojson, f, indent=2)
        return out_file

    def export_kml(self, out_file: Path) -> Path:
        out_file.parent.mkdir(parents=True, exist_ok=True)
        kml = f"""<?xml version="1.0" encoding="UTF-8"?>
<kml xmlns="http://www.opengis.net/kml/2.2">
  <Document>
    <name>{self.case.case_id}</name>
  </Document>
</kml>"""
        with open(out_file, "w", encoding="utf-8") as f:
            f.write(kml)
        return out_file

    def export_xlsx(self, out_file: Path) -> Path:
        try:
            import openpyxl
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Summary"
            ws.append(["Case ID", self.case.case_id])
            ws.append(["Case Name", self._filter_str(self.data.get("case_name", ""))])
            ws.append(["Analyst", self.data.get("analyst", "")])

            ws_ioc = wb.create_sheet(title="IOCs")
            ws_ioc.append(["ID", "Type", "Value", "Confidence"])
            for i in self.data.get("iocs", []):
                ws_ioc.append([i.get("id"), i.get("type"), self._filter_str(i.get("value")), i.get("confidence")])

            wb.save(out_file)
            return out_file
        except ImportError:
            raise

    def export_docx(self, out_file: Path) -> Path:
        try:
            import docx
            doc = docx.Document()
            doc.add_heading(f"TraceForge Report — {self.case.case_id}", 0)
            doc.add_paragraph(f"Case: {self._filter_str(self.data.get('case_name',''))}")
            doc.add_paragraph(f"Analyst: {self.data.get('analyst','')}")
            doc.save(out_file)
            return out_file
        except ImportError:
            raise
